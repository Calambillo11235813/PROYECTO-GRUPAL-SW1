import os
from typing import Tuple
import PyPDF2
from docx import Document
from django.core.files.uploadedfile import UploadedFile

class FileProcessor:
    """Procesador de archivos para extracción de texto - Versión estable"""
    
    ALLOWED_EXTENSIONS = ['.txt', '.pdf', '.docx']
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @staticmethod
    def validate_file(file: UploadedFile) -> Tuple[bool, str]:
        """Valida el archivo subido"""
        # Validar tamaño
        if file.size > FileProcessor.MAX_FILE_SIZE:
            return False, "El archivo es demasiado grande (máximo 10MB)"
        
        # Validar extensión
        file_extension = os.path.splitext(file.name)[1].lower()
        if file_extension not in FileProcessor.ALLOWED_EXTENSIONS:
            return False, f"Formato no soportado. Use: {', '.join(FileProcessor.ALLOWED_EXTENSIONS)}"
        
        # Validación básica de contenido por extensión
        try:
            file.seek(0)
            first_bytes = file.read(8)
            file.seek(0)
            
            # Validaciones básicas por tipo
            if file_extension == '.pdf' and not first_bytes.startswith(b'%PDF'):
                return False, "El archivo no es un PDF válido"
            elif file_extension == '.docx' and not (b'PK' in first_bytes or b'zip' in str(first_bytes).lower()):
                # DOCX son archivos ZIP
                pass  # Validación más permisiva para DOCX
                
        except Exception:
            pass  # Si hay error en validación, continuar
        
        return True, "Archivo válido"
    
    @staticmethod
    def extract_text(file: UploadedFile) -> Tuple[bool, str]:
        """Extrae texto del archivo"""
        try:
            file_extension = os.path.splitext(file.name)[1].lower()
            
            if file_extension == '.txt':
                return FileProcessor._extract_from_txt(file)
            elif file_extension == '.pdf':
                return FileProcessor._extract_from_pdf(file)
            elif file_extension == '.docx':
                return FileProcessor._extract_from_docx(file)
            else:
                return False, "Formato no soportado"
                
        except Exception as e:
            return False, f"Error al procesar archivo: {str(e)}"
    
    @staticmethod
    def _extract_from_txt(file: UploadedFile) -> Tuple[bool, str]:
        """Extrae texto de archivo TXT"""
        try:
            # Intentar diferentes codificaciones comunes
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    content = file.read().decode(encoding)
                    if content.strip():
                        return True, content.strip()
                except UnicodeDecodeError:
                    continue
            
            return False, "No se pudo decodificar el archivo de texto"
            
        except Exception as e:
            return False, f"Error al leer archivo TXT: {str(e)}"
    
    @staticmethod
    def _extract_from_pdf(file: UploadedFile) -> Tuple[bool, str]:
        """Extrae texto de archivo PDF"""
        try:
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
            
            if len(pdf_reader.pages) == 0:
                return False, "El PDF está vacío"
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content.append(page_text.strip())
                except Exception as e:
                    print(f"Error extrayendo página {page_num}: {e}")
                    continue
            
            if not text_content:
                return False, "No se pudo extraer texto del PDF (puede estar escaneado)"
            
            full_text = '\n\n'.join(text_content)
            return True, full_text
            
        except Exception as e:
            return False, f"Error al leer archivo PDF: {str(e)}"
    
    @staticmethod
    def _extract_from_docx(file: UploadedFile) -> Tuple[bool, str]:
        """Extrae texto de archivo DOCX"""
        try:
            file.seek(0)
            doc = Document(file)
            
            text_content = []
            
            # Extraer párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            text_content.append(cell.text.strip())
            
            if not text_content:
                return False, "El documento Word está vacío"
            
            full_text = '\n\n'.join(text_content)
            return True, full_text
            
        except Exception as e:
            return False, f"Error al leer archivo DOCX: {str(e)}"